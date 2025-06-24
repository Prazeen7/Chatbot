import re
import os
import pytesseract
from PIL import Image
import io
from pypdf import PdfReader
from docx import Document
from warnings import warn
from typing import List
import tabula
from config import Config

def validate_file(file_path: str):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at {file_path}")
    if not (file_path.lower().endswith('.pdf') or file_path.lower().endswith('.docx')):
        raise ValueError(f"Unsupported file format: {file_path}. Only PDF and DOCX are supported.")

def check_ocr_capability() -> bool:
    try:
        import pytesseract
        return True
    except ImportError:
        warn("OCR dependencies not found. Image processing disabled.")
        return False

def split_into_chunks(text: str) -> List[str]:
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        if len(current_chunk) + len(para) > Config.MAX_CHUNK_SIZE:
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = para[-Config.CHUNK_OVERLAP:] + " " if Config.CHUNK_OVERLAP else ""
            else:
                chunks.append(para[:Config.MAX_CHUNK_SIZE])
                current_chunk = para[Config.MAX_CHUNK_SIZE - Config.CHUNK_OVERLAP:] + " "
        else:
            current_chunk += " " + para if current_chunk else para
    
    if current_chunk:
        chunks.append(current_chunk)
        
    return [chunk for chunk in chunks if len(chunk) >= Config.MIN_CHUNK_SIZE]

def process_images(page, can_ocr: bool) -> List[str]:
    if not can_ocr:
        return []
    
    img_texts = []
    for image_file in page.images:
        try:
            img = Image.open(io.BytesIO(image_file.data))
            if img.mode == 'RGBA':
                img = img.convert('RGB')
            img_text = pytesseract.image_to_string(img)
            if img_text.strip():
                img_texts.append(img_text.strip())
        except Exception as e:
            warn(f"Image processing skipped: {str(e)[:100]}...")
    return img_texts

def extract_text_from_pdf(pdf_path: str, can_ocr: bool) -> List[str]:
    text_chunks = []
    try:
        reader = PdfReader(pdf_path)
        
        # Extract text and tables from each page
        for page_num, page in enumerate(reader.pages, 1):
            # Extract regular text
            if text := page.extract_text():
                text_chunks.extend(split_into_chunks(text))
            
            # Extract images (if OCR enabled)
            if can_ocr:
                text_chunks.extend(process_images(page, can_ocr))
            
            # Extract tables
            try:
                tables = tabula.read_pdf(pdf_path, pages=page_num, multiple_tables=True)
                for table in tables:
                    # Convert table to string (join rows with newlines)
                    table_text = '\n'.join(table.fillna('').astype(str).apply(lambda x: ' | '.join(x), axis=1))
                    if table_text.strip():
                        text_chunks.extend(split_into_chunks(table_text))
            except Exception as e:
                warn(f"Table extraction error for {pdf_path} page {page_num}: {str(e)}")
    
    except Exception as e:
        warn(f"PDF processing error for {pdf_path}: {str(e)}")
    
    return text_chunks

def extract_text_from_docx(docx_path: str) -> List[str]:
    text_chunks = []
    try:
        doc = Document(docx_path)
        full_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
        
        if full_text:
            text_chunks.extend(split_into_chunks('\n'.join(full_text)))
    
    except Exception as e:
        warn(f"DOCX processing error for {docx_path}: {str(e)}")
    
    return text_chunks

def save_chunks_to_file(chunks: List[str], output_file: str) -> str:
    """Save text chunks to a DOCX file, creating directories if needed."""
    try:
        # Ensure output_file is valid
        if not output_file:
            output_file = os.path.join("output", "chunks.docx")
        
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(output_file) or '.', exist_ok=True)
        
        # Create a new DOCX document
        doc = Document()
        doc.add_heading('Text Chunks', level=1)
        
        for i, chunk in enumerate(chunks, 1):
            doc.add_heading(f'Chunk {i}', level=2)
            doc.add_paragraph(chunk)
            doc.add_paragraph()  # Add empty line for separation
        
        doc.save(output_file)
        print(f"Saved {len(chunks)} chunks to: {os.path.abspath(output_file)}")
        return output_file
    except Exception as e:
        warn(f"Failed to save chunks: {str(e)}")
        return ""

def load_documents(save_chunks: bool = False, output_file: str = "") -> List[str]:
    """Extract text from all PDF and DOCX files in the data folder, optionally save chunks to file."""
    data_folder = Config.DATA_FOLDER
    if not os.path.exists(data_folder):
        raise FileNotFoundError(f"Data folder not found at {data_folder}")
    
    file_paths = [
        os.path.join(data_folder, f) for f in os.listdir(data_folder)
        if f.lower().endswith(('.pdf', '.docx'))
    ]
    
    if not file_paths:
        raise ValueError("No PDF or DOCX files found in the data folder")
    
    for file_path in file_paths:
        validate_file(file_path)
    
    can_ocr = check_ocr_capability()
    text_chunks = []
    
    for file_path in file_paths:
        try:
            if file_path.lower().endswith('.pdf'):
                text_chunks.extend(extract_text_from_pdf(file_path, can_ocr))
            elif file_path.lower().endswith('.docx'):
                text_chunks.extend(extract_text_from_docx(file_path))
        except Exception as e:
            warn(f"Error processing {file_path}: {str(e)}")
            continue
    
    filtered_chunks = [chunk for chunk in text_chunks if chunk]
    
    if save_chunks:
        saved_path = save_chunks_to_file(filtered_chunks, output_file)
        if not saved_path:
            warn("Chunks were processed but not saved to file")
    
    return filtered_chunks